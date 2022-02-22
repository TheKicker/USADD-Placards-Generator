
# From the documentation: 
# https://python-docx.readthedocs.io/en/latest/

from docx import Document

Document = Document()

Document.add_heading('Document Title', 0)

p = Document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

Document.add_heading('Heading, level 1', level=1)

Document.add_page_break()

Document.save('demo.docx')