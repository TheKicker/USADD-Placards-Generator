import docx, os, csv, urllib.request, qrcode
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from PIL import Image 

# Fetch the current working directory of the script
folder = os.getcwd()

# Function to concatenate the images
# https://note.nkmk.me/en/python-pillow-concat-images/#concatenate-multiple-images-at-once
def get_concat_h_multi_resize(im_list, resample=Image.BICUBIC):
            min_height = min(im.height for im in im_list)
            im_list_resize = [im.resize((int(im.width * min_height / im.height), min_height),resample=resample)
                            for im in im_list]
            total_width = sum(im.width for im in im_list_resize)
            dst = Image.new('RGB', (total_width, min_height))
            pos_x = 0
            for im in im_list_resize:
                dst.paste(im, (pos_x, 0))
                pos_x += im.width
            return dst


# Open the CSV file and read it
with open('products_export.csv', 'r') as csv_file:
    csv_reader = csv.reader(csv_file)
    i = 0
    # Skips row 1 (the column titles)
    next(csv_reader)

    # Read the CSV row by row
    for line in csv_reader:
        i = i + 1
        # Create a new document
        doc = docx.Document()
        section = doc.sections[-1]

        # Rotate the document to landscape mode
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height

        # Read the CSV column by column
        title = line[0]
        sku = line[1]
        url = line[2]

        # Shopify sometimes adds a ' before SKUs, this ignores if exists
        if "'" in sku:
            sku = sku[1:]
            print("' Detected, skipping first char in {sku1}.  Document {page} complete. ".format(sku1 = sku, page = i))
        else:
            print("{sku1} ready, document {page} complete. ".format(sku1 = sku, page = i))
            pass

        # Change directory to images folder
        os.chdir(folder + "/images")

        # Fetch the image from the URL
        r = urllib.request.urlopen(url)

        # Check the ending on the image
        if ".png?" in url:
            suffix = ".png"
        elif ".jpg?" in url:
            suffix = ".jpg"
        else:
            suffix = ".jpeg"

        # Save the image
        with open(sku + suffix, "wb") as f:
            f.write(r.read())

        # Change directory to QR Code folder
        os.chdir(folder + "/images/qr-codes") 

        # Generate the QR Code
        img = qrcode.make(sku)
        img.save(sku + "-qrcode.png")

        # Change directory to original location
        os.chdir(folder)

        # SKU line (the sku of the product) & styling
        sku_head = doc.add_heading(sku, 0)
        sku_head.style = doc.styles.add_style('Style Name SKU', WD_STYLE_TYPE.PARAGRAPH)
        sku_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font = sku_head.style.font
        font.name = 'Times New Roman'
        font.size = Pt(85)
        font.bold = True
        font.italic = False

        # Title line (the name of the product) & styling
        title_head = doc.add_heading(title, 1)
        title_head.style = doc.styles.add_style('Style Name TITLE', WD_STYLE_TYPE.PARAGRAPH)
        title_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font = title_head.style.font
        font.name = 'Calibri'
        font.size = Pt(34)
        font.color.rgb = RGBColor(0,0,0)
        font.bold = False
        title_head = title_head.paragraph_format
        title_head.space_after = Pt(40)

        # Populate image variables
        im1 = Image.open(folder + "/images/" + sku + suffix)
        im2 = Image.open(folder + "/images/qr-codes/" + sku + "-qrcode.png")
        im3 = Image.open("assets/usadd-logo.png")

        # Concatenate the images and save to a folder
        get_concat_h_multi_resize([im1, im2, im3]).save(folder + "/images/concat/" + sku + '-resize.jpg')
        
        # Add the picture to the document
        doc.add_picture(folder + "/images/concat/" + sku + "-resize.jpg", width=Inches(9))

        # Save the document to the output folder
        doc.save(folder + "/output/" + sku + ".docx")

    
    print(" Job finished with {pages} complete. ".format(pages=i))
